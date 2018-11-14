# encoding=utf-8
# Program: 这是一个实现输出回测结果的脚本程序
# Writer: Qijie Li
# Date: Oct.21.2017

import copy
import os
import numpy as np
import pandas as pd
import shutil
import matplotlib as mpl
import matplotlib.pyplot as plt
from pylab import text
from docx import Document
from docx.shared import Inches
from datetime import date
import RecordOutput
mpl.rcParams['font.sans-serif'] = ['SimHei'] #用来正常显示中文标签
mpl.rcParams['axes.unicode_minus'] = False      # 解决保存图像是负号'-'显示为方块的问题
font = {'family': 'serif',
         'style': 'normal',
         'weight': 'normal',
        'color':  'black',
        'size': 12
        }

# 静态等权的类
__mataclass__=type # Use new type classes
class OutputClass:
    '''
    输出结果类
    功能：画图，输出word报告，生成excel交易记录
    公共方法：
    plotPortfolioRet（）: 画出组合收益
    WinRate（）：胜率
    plotNetVal（）：画出净值曲线
    plotPeriodRet （）： 画出年度收益图
    Report_multipleFactors() ：返回回测报告
    writeHistoryRecord()
    '''

    # <=========================================基本框架=========================================>
    def __init__(self, inputData):
        '''
        设置控制变量： 调仓周期、调仓类型、分组数、开仓金额、手续费等
        '''
        # 设置参数
        self.inputData = copy.deepcopy(inputData)
        self.timeFreq = self.inputData[u'timeFreq']
        self.dateFreq = self.inputData[u'dateFreq']
        self.fullIndx = pd.Index(self.dateFreq)
        if self.inputData[u'benchmark']:
            self.benchmarkName = self.inputData[u'benchmark'].keys()
            self.benchmark = self.inputData[u'benchmark'][self.benchmarkName[0]][u'close']
            self.benchmarkRet = self.benchmark/self.benchmark.shift(1).fillna(method='bfill') - 1
            self.benchmarkNetValue = 1 + self.benchmarkRet.cumsum()
        else:
            self.benchmarkName = None
            self.benchmark, self.benchmarkRet, self.benchmarkNetValue = [pd.Series()]*3
        self.strategyNameList = self.inputData[u'strategyNameList']
        # 格式转换
        tempDicFund = \
            {i: (self.inputData[u'outPutData'][i][u'fundRecord']) for i in self.strategyNameList}
        self.initiCapital = {i: self.inputData[u'outPutData'][i][u'intiFund'] for i in self.strategyNameList}
        self.fundCumsum = \
            {i: self.__dictToSeries(tempDicFund[i]) for i in self.strategyNameList}             # 累计资金曲线
        self.portValue = \
            {i: (self.fundCumsum[i]/self.inputData[u'outPutData'][i][u'intiFund']) for i in self.strategyNameList}
        if self.inputData[u'benchmark']:
            self.portExValue = \
                {i: 1 + (self.portValue[i] - self.benchmarkNetValue) for i in self.strategyNameList}# 超额收益率曲线
        else:
            self.portExValue = None
        self.orderRecord = \
            {i: self.inputData[u'outPutData'][i][u'orderRecord'] for i in self.strategyNameList}
        self.PLRecord = \
            {i: self.inputData[u'outPutData'][i][u'PLRecord'] for i in self.strategyNameList}
        self.positionRecord = \
            {i: self.inputData[u'outPutData'][i][u'positionRecord'] for i in self.strategyNameList}
        self.outPutData = {}
        self.outPutData[u'strategyDict'] = {}
        self.today = date.today().isoformat()
        # 画图路径
        self.plotPath = u'../Output/plot_results'
        self.benchmarkNameDict = {u'wind_a_stocks': u'万得全A', u'hs300_stocks': u'沪深300',
                                  u'zz500_stocks': u'中证500', u'sz50_stocks': u'上证50'}

    def __dictToSeries(self, dictData):
        tempData = [dictData[i] for i in dictData]
        locTimeIndex = [i for i in dictData]
        series = pd.Series(tempData, index=locTimeIndex)[self.fullIndx]
        return series

    def __fundToRet(self, fund):
        ret = fund/fund.shift(1).fillna(method='bfill') - 1
        return ret

    def adjustDateFreq(self, transferCycle, data):
        '''
        时间频率转换函数：如周调仓，将导入的原始数据都转化为周频率数据
        '''
        data0 = copy.deepcopy(data)
        if transferCycle in [u"W", u"M", u"A"]: # 周度,月度, 年度调仓
            tempIndex = pd.to_datetime(data0.index) # index转化为时间格式
            index_new = tempIndex.to_period(freq = transferCycle).to_series().astype(str)
            uniqueFreq = np.unique(index_new) # 保留不重复得日期
            data0.index = index_new
        elif type(transferCycle) == int:  # 按X天为周期调仓
            originalIndex = data0.index
            index = originalIndex.to_series().astype(str)
            index = index.tolist()
            p = transferCycle
            temp1 = originalIndex[::p] # 取周期头数据
            temp2 = originalIndex[p - 1:][::p] # 取周期尾数据
            indexStart = temp1.to_series().astype(str)
            indexStart = indexStart.tolist()
            indexEnd = temp2.to_series().astype(str)
            indexEnd = indexEnd.tolist()
            if len(indexStart) > len(indexEnd): # 处理周期头长度大于周期末长度
                indexEnd.append(index[-1])
            final = [x + r"/" + y for x, y in zip(indexStart, indexEnd)]
            final = pd.Series(final, index = final)
            tempIndex = [0] * len(originalIndex) # 不重复的index
            for i in range(len(final) - 1): # 遍历不同周期，每个相同周期, 创建相同的Index
                temp = final[i]
                tempIndex[i*p : (i+1)*p] = [temp]*p  # 创建新的index来覆盖掉原来的index
            for j in range((len(final) - 1) * p, len(tempIndex)): # 处理最后一个周期的index
                tempIndex[j] = final[-1]
            index_new = pd.Series(tempIndex)
            uniqueFreq = final  # 保留不重复得日期
            data0.index = index_new
        else:
            print(u"请输入正确调仓时段格式")
        return data0, uniqueFreq

    def plotNetVal(self, excessRet=False):
        '''
        # todo 重新封装画图函数
        :param excessRet: 输出超额资金收益.csv
        :return:         择时前后的净值图曲线（单利）
        '''
        if not self.benchmarkNetValue.empty:
            benchmarkValue = copy.deepcopy(self.benchmarkNetValue.values)
            benchmarkValue.shape = (benchmarkValue.shape[0], 1)
            exDailyValue = copy.deepcopy(self.portExValue)
        else:
            benchmarkValue = None
        fundCumsum = copy.deepcopy(self.fundCumsum)
        dailyValue = copy.deepcopy(self.portValue)
        timeFreq = copy.deepcopy(self.dateFreq)
        colorOption = [u'orangered', u"steelblue", u"violet", u"seagreen", u"crimson", u"fuchsia"]
        strategyColor = {self.strategyNameList[i]: colorOption[i]
                         for i in range(len(self.strategyNameList))}
        self.outPutData[u'strategyDict'][u'drawdown'] = {}
        self.outPutData[u'strategyDict'][u'finalValue'] = {}
        # <=========================================收益曲线=========================================>
        # 绘制净值曲线图
        # 资金曲线
        size = (16, 12)
        fig1 = plt.figure(figsize=(size[0], size[1]))
        plt.subplots_adjust(hspace=0)
        for key in self.strategyNameList:
            x = fundCumsum[key]
            x = x.values
            plt.plot(x)
            # 输出超额资金收益.csv
            if excessRet:
                outPutDF = pd.DataFrame(
                    exDailyValue[key], columns=[u'超额收益'])
                outPutDF.to_csv(
                    '../Output/' + self.strategyNameList[0] + ' ex_return' + self.today + '.csv', encoding='gbk')
        xtickActuallyLen = 0.4
        gapNum = size[0]/xtickActuallyLen
        gap = max(1, int(round(len(timeFreq)/float(gapNum))))
        plt.xticks(range(0, len(timeFreq), gap), timeFreq[::gap], rotation=45, fontsize=12, weight='bold')  # x轴刻度
        fig1.savefig(self.plotPath + u'/fund_daily.jpg')

        # 收益曲线
        fig2 = plt.figure(figsize=(size[0], size[1]))
        fig2.set_size_inches(15, 10, forward=True)
        plt.subplots_adjust(hspace=0)
        if not self.benchmarkNetValue.empty:
            ax1 = fig2.add_subplot(211)
            ax2 = plt.subplot(212)
            if self.benchmarkName[0] in self.benchmarkNameDict.keys():
                plotLabel = self.benchmarkNameDict[self.benchmarkName[0]]
            else:
                plotLabel = self.benchmarkName[0].strip(u'_stocks')
            ax1.plot(benchmarkValue, label=plotLabel, color='seagreen')
            ax1.annotate(u"总净值:" + str('%.3f' % benchmarkValue[-1][0]),
                         xy=(len(benchmarkValue) - 1, benchmarkValue[-1][0]),
                         xytext=(len(benchmarkValue) - 1, benchmarkValue[-1][0] + 0.01), fontsize=12)
            yLim2 = ax2.get_ylim()
        else:
            ax1 = fig2.add_subplot(111)
        ax1.set_title(u'资金日线收益曲线', fontsize=24, weight='bold')
        ax1.set_ylabel(u'累计绝对收益', fontsize=20)

        yLim1 = ax1.get_ylim()
        for key in self.strategyNameList:
            length = len(timeFreq)
            x = dailyValue[key].values
            x.shape = (x.shape[0], 1)
            if not self.benchmarkNetValue.empty:
                exValue = exDailyValue[key].values
                exValue.shape = (exValue.shape[0], 1)
            # <=========================================回撤程序=========================================>
            drawdown, exDrawdown, = np.zeros(length), np.zeros(length)  # drawdown_timing = np.zeros(length)
            for i in range(length):
                dd = (x[:1 + i].max() - x[i])/x[:1 + i].max()
                drawdown[i] = dd
                if not self.benchmarkNetValue.empty:
                    excess_dd = (exValue[:1 + i].max() - exValue[i])/exValue[:1 + i].max()
                    exDrawdown[i] = excess_dd
            maxDd = drawdown.max()                                      # 绝对收益率的最大回撤
            maxDd_index = drawdown.argmax()
            maximum = x[:maxDd_index + 1].argmax()                      # 绝对收益率最大的index
            dot_ind = int((maxDd_index + maximum) / 2)                  # 绝对收益率的最大回撤指向位置
            dot_y = 0.5 * x[maximum][0] + 0.5 * x[maxDd_index][0]
            maxDd_duration = maxDd_index - maximum                      # 绝对收益率的最大回撤持续期
            if not self.benchmarkNetValue.empty:
                maxExDd = exDrawdown.max()                                  # 超额收益率的最大回撤
                maxExDd_index = exDrawdown.argmax()
                maximumExcess = exValue[:maxExDd_index + 1].argmax()        # 超额收益率最大的index
                excess_dot_ind = int((maxExDd_index + maximumExcess) / 2)   # 超额收益率的最大回撤指向位置
                excess_dd_y = 0.5 * exValue[maximumExcess][0] + 0.5 * exValue[maxExDd_index][0]
                maxExDd_duration = maxExDd_index - maximumExcess            # 超额收益率的最大回撤持续期
            else:
                maxExDd = None
                maxExDd_duration = None
            # <=========================================回撤程序=========================================>
            ax1.plot(x, label=key, color='red')
            ax1.plot(maxDd_index, x[maxDd_index], u'p', color=u'Black')
            ax1.axvline(maxDd_index, ls='--', color=u'dimgrey')
            ax1.plot(maximum, x[maximum], 'p', color=u'Black')
            ax1.axvline(maximum, ls='--', color=u'dimgrey')
            # annotate注释
            if not self.benchmarkNetValue.empty:
                graphLineGap = (yLim2[1] - yLim2[0]) * 0.5
            else:
                graphLineGap = (yLim1[1] - yLim1[0]) * 0.5

            ax1.annotate(u'最大回撤率: ' + str(round(maxDd * 100., 2)) + u'%' +
                         u'\n' + u"最大回撤持续期:" + str(maxDd_duration) + u'天',
                         xy=(dot_ind, dot_y + graphLineGap * 0.01),
                         xytext=(dot_ind + round(0.1 * length), dot_y + graphLineGap * 0.1),
                         color=u'black', fontsize=12,
                         arrowprops=dict(facecolor=u'orange', arrowstyle='->'),
                         horizontalalignment=u'left', verticalalignment=u'top')

            ax1.annotate(u"总净值:" + str(u'%.3f' % x[-1][0]), xy=(len(x) -1, x[-1][0]),
                         xytext=(len(x) - 1, x[-1][0] + 0.01), color=u'black', fontsize=12)
            if not self.benchmarkNetValue.empty:
                graphLineGap2 = (yLim1[1] - yLim1[0]) * 0.5
                # 图二
                ax2.plot(exValue, label=u'超额收益', color='orangered')
                ax2.plot(maxExDd_index, exValue[maxExDd_index], u'p', color=u'Black')
                plt.xticks(range(0, len(timeFreq), gap), timeFreq[::gap], rotation=45, fontsize=12) # x轴刻度
                ax2.axvline(maxExDd_index, ls='--', color=u'dimgrey')
                ax2.plot(maximumExcess, exValue[maximumExcess], u'p', color=u'Black')
                ax2.axvline(maximumExcess, ls='--', color=u'dimgrey')
                # 注释
                ax2.annotate(u'超额收益最大回撤率: ' + str(round(maxExDd * 100, 2)) + u'%' +
                             u'\n' + u"超额收益最大回撤持续期:" + str(maxExDd_duration) + u'天',
                             xy=(excess_dot_ind, excess_dd_y + graphLineGap2 * 0.01),
                             xytext=(excess_dot_ind + round(0.1 * length), excess_dd_y + graphLineGap2 * 0.1),
                             color=u'black', fontsize=12,
                             arrowprops = dict(facecolor=u'orange', arrowstyle='->'),
                             horizontalalignment=u'left', verticalalignment=u'top')

                ax2.annotate(u"总净值:" + str(u'%.3f' % exValue[-1][0]),
                             xy=(len(exValue) - 1, exValue[-1][0]),
                             xytext=(len(exValue) - 1, exValue[-1][0] + 0.01),
                             color=u'black', fontsize=12)

            plt.xticks(range(0, len(timeFreq), gap), timeFreq[::gap], rotation=45, fontsize=12, weight='bold')  # x轴刻度
            self.outPutData[u'strategyDict'][u'drawdown'][key] = \
                {u'maxDrawdown': maxDd,
                 u'maxDrawdown_duration': maxDd_duration}
            self.outPutData[u'strategyDict'][u'finalValue'][key] = \
                {u'finalAbsValue': str(u'%.3f' % x[-1][0])}
            if not self.benchmarkNetValue.empty:
                self.outPutData[u'strategyDict'][u'drawdown'][key][u'maxExcessDrawdown'] = maxExDd
                self.outPutData[u'strategyDict'][u'drawdown'][key][u'maxExcessDrawdown_duration'] = maxExDd_duration
                self.outPutData[u'strategyDict'][u'finalValue'][key][u'finalExcessValue'] = str(u'%.3f' % exValue[-1][0])
        pass   # 策略循环结束

        for tempAx in [ax1, ax2]:
            for tick in tempAx.yaxis.get_major_ticks():        # 设置轴字体大小
                tick.label1.set_fontsize(12)

        if not self.benchmarkNetValue.empty:
            ax2.set_ylabel(u'超额收益', fontsize=20)
            ax1.legend(loc=u'lower center', fontsize=16)
            ax2.legend(loc=u'lower center', fontsize=16)
        else:
            ax1.set_ylabel(u'累计绝对收益', fontsize=16)
            ax1.legend(loc=u'lower center', fontsize=12)
        fig2.savefig(self.plotPath + u'/netValue_daily.jpg')
        print('资金曲线画图完成')

    def periodRet(self):
        '''
        时段收益率计算程序+画图
        注意事项：夏普率，信息比率时月度级别
        :return: 夏普率，信息比率，年度收益画图
        '''
        k = int(1)
        cycleSet = {u'月度': u'M', u'年度': u'A'}
        portValue = copy.deepcopy(self.portValue)
        if self.benchmarkNetValue:
            portExValue = copy.deepcopy(self.portExValue)
        self.outPutData[u'strategyDict'][u'sharpRatio'] = {}
        self.outPutData[u'strategyDict'][u'informationRatio'] = {}
        fig2 = plt.figure(figsize=(16, 9))
        for key in cycleSet:
            data = portValue
            excessData = portExValue
            for w in data:  # w--不同策略
                # 计算X度收益率
                tempRet = pd.DataFrame(data[w], index=timeFreq)
                tempExRet = pd.DataFrame(excessData[w], index=timeFreq)
                # 绝对收益改变为X度标签
                tempPeriodRet, periodIndex = self.adjustDateFreq(cycleSet[key], tempRet)
                tempPeriodExRet = self.adjustDateFreq(cycleSet[key], tempExRet)[0]
                # 超额收益改变为X度标签
                periodRet = pd.DataFrame(
                    np.zeros(len(periodIndex)), index=periodIndex, columns=[key + u'收益'])  # 月度绝对收益
                periodExRet = copy.deepcopy(periodRet)  # X度超额收益
                for i in range(len(periodIndex)):
                    if not i:
                        if len(tempPeriodRet.loc[periodIndex[i]].shape) > 1:
                            periodRet.iloc[i] = tempPeriodRet.loc[periodIndex[i]].tail(1).values - 1
                            periodExRet.iloc[i] = tempPeriodExRet.loc[periodIndex[i]].tail(1).values
                        else:  # 开始日期只有在X期末
                            periodRet.iloc[i] = tempPeriodRet.loc[periodIndex[i]]
                            periodExRet.iloc[i] = tempPeriodExRet.loc[periodIndex[i]]
                    else:
                        periodRet.iloc[i] = tempPeriodRet.loc[periodIndex[i]].tail(1).values - \
                                           tempPeriodRet.loc[periodIndex[i - 1]].tail(1).values
                        periodExRet.iloc[i] = tempPeriodExRet.loc[periodIndex[i]].tail(1).values - \
                                             tempPeriodExRet.loc[periodIndex[i - 1]].tail(1).values
                pass
                if key == u'月度':
                    sharp_ratio = np.mean(periodRet) / np.std(periodRet)
                    IR = np.mean(periodExRet) / np.std(periodExRet)
                    # 月度夏普率和月度信息比率
                    self.outPutData[u'strategyDict'][u'sharpRatio'][w] = sharp_ratio.values[0]
                    self.outPutData[u'strategyDict'][u'informationRatio'][w] = IR.values[0]
                if key == u'年度':
                    # 绝对年度收益柱状图
                    ax = fig2.add_subplot(2, len(data), k)
                    ax.set_title(w)
                    ax.set_ylabel(u'年度绝对收益率')
                    listRet = periodRet.iloc[:, 0].tolist()
                    ax.bar(range(len(periodIndex)), listRet, color=u'coral')
                    # 标注
                    yLim1 = ax.get_ylim()
                    ax.annotate(str(periodIndex[-1]) + u'年年度收益率 \n' + str(periodRet.iloc[-1].values[-1]),
                                 xy=(len(periodIndex) - 1, periodRet.iloc[-1].values[-1]),
                                xytext=(len(periodIndex) - 1, periodRet.iloc[-1].values[-1]
                                        + (yLim1[1] - yLim1[0]) * 0.05),
                                 color=u'black', fontsize=10, verticalalignment=u'top')
                    # 超额年度收益柱状图
                    ax = fig2.add_subplot(2, len(data), k + + len(data))
                    ax.set_title(w)
                    ax.set_xlabel(u'年')
                    ax.set_ylabel(u'年度超额收益率')
                    listExRet = periodExRet.iloc[:, 0].tolist()
                    ax.bar(range(len(periodIndex)), listExRet, color=u'coral')
                    # 标注
                    yLim2 = ax.get_ylim()
                    ax.annotate(str(periodIndex[-1]) + u'年年度收益率 \n' + str(periodExRet.iloc[-1].values[-1]),
                                xy=(len(periodIndex) - 1, periodExRet.iloc[-1].values[-1]),
                                xytext=(len(periodIndex) - 1, periodExRet.iloc[-1].values[-1]
                                        + (yLim2[1] - yLim2[0]) * 0.05),
                                color=u'black', fontsize=10, verticalalignment=u'top')
                    plt.xticks(np.arange(len(periodIndex)), periodIndex, rotation=45)  # x轴刻度
                    k += 1
                pass  # 单个策略结束
            pass # 单个周期结束
        fig2.savefig(self.plotPath + u'/annualReturns.jpg')
        print('时段函数计算结束')
        # end of the function

    def recordOutput(self):
        print('输出成交记录')
        for key in self.strategyNameList:
            orderDict = self.inputData[u'outPutData'][key]
            RecordOutput.transitionWriter(self.timeFreq, self.dateFreq, key, orderDict)
            # RecordOutput.evaluation_F()
        print('成交记录输出完毕')
        print('评价指标生成完毕')
    # <=========================================写入报告=========================================>
    def Report_multipleFactors(self, tester, testContent):
        '''
        将运行结果写入报告
        '''
        strategyNamelist = copy.deepcopy(self.strategyNameList)
        document = Document()
        document.add_heading(u'策略回测报告', 0)
        document.add_paragraph(u'')
        document.add_paragraph(u'动量策略回测结果如下')
        # 设置表格
        table1 = document.add_table(rows=16, cols=2, style=u'Table Grid')
        i = 0
        table1.cell(i, 0).text = u'测试内容'
        table1.cell(i, 1).text = testContent
        i += 1
        table1.cell(i, 0).text = u'测试人'
        table1.cell(i, 1).text = tester
        i += 1
        table1.cell(i, 0).text = u'测试时间'
        table1.cell(i, 1).text = self.today
        i += 1
        table1.cell(i, 0).text = u'样本选取时间'
        table1.cell(i, 1).text = self.timeFreq[0] + u'至' + self.timeFreq[-1]
        i += 1
        if self.benchmarkName:
            table1.cell(i, 0).text = u'基准市场'
            table1.cell(i, 1).text = self.benchmarkName
            i += 1
        table1.cell(i, 0).text = u'股票池'
        table1.cell(i, 1).text = u'全市场'
        i += 1
        table1.cell(i, 0).text = u'手续费'
        table1.cell(i, 1).text = str(self.inputData[u'fee'])
        i += 1
        # table1.cell(i, 0).text = u'绝对胜率'
        # table1.cell(i, 1).text = "\n".join([(j + ":" + str(
        #     self.outPutData[u'strategyDict'][u'winRate'][j][u'absWinRate'])
        #                                      ) for j in strategyNamelist])
        # i += 1
        # table1.cell(i, 0).text = u'超额胜率'
        # table1.cell(i, 1).text = "\n".join([(j + ":" + str(
        #     self.outPutData[u'strategyDict'][u'winRate'][j][u'excessWinRate'])
        #                                      ) for j in strategyNamelist])
        # i += 1
        table1.cell(i, 0).text = u'最大回撤率'
        table1.cell(i, 1).text = "\n".join([(j+":" + str(
            self.outPutData[u'strategyDict'][u'drawdown'][j][u'maxDrawdown'])
                                             ) for j in strategyNamelist])
        i += 1
        table1.cell(i, 0).text = u'最大回撤持续期'
        table1.cell(i, 1).text = "\n".join([(j+":" + str(
            self.outPutData[u'strategyDict'][u'drawdown'][j][u'maxDrawdown_duration'])
                                             + u'天') for j in strategyNamelist])
        i += 1
        if self.benchmarkName:
            table1.cell(i, 0).text = u'超额最大回撤率'
            table1.cell(i, 1).text = "\n".join([(j+":" + str(
                self.outPutData[u'strategyDict'][u'drawdown'][j][u'maxExcessDrawdown'])
                                                 ) for j in strategyNamelist])
            i += 1
            table1.cell(i, 0).text = u'超额最大回撤持续期'
            table1.cell(i, 1).text = "\n".join([(j+":" + str(
                self.outPutData[u'strategyDict'][u'drawdown'][j][u'maxExcessDrawdown_duration'])
                                                 + u'天') for j in strategyNamelist])
        i += 1
        if u'sharpRatio' in self.outPutData[u'strategyDict'].keys():
            table1.cell(i, 0).text = u'夏普比率'
            table1.cell(i, 1).text = "\n".join([(j+":" + str(
                self.outPutData[u'strategyDict'][u'sharpRatio'][j])
                                                 ) for j in strategyNamelist])
            i += 1
        if u'informationRatio' in self.outPutData[u'strategyDict'].keys():
            table1.cell(i, 0).text = u'信息比率'
            table1.cell(i, 1).text = "\n".join([(j+":" + str(
                self.outPutData[u'strategyDict'][u'informationRatio'][j])
                                                 ) for j in strategyNamelist])

        document.add_paragraph("")
        # 插入图片
        document.add_paragraph(u'资金盈亏如下').paragraph_format.keep_with_next
        document.add_picture(self.plotPath + u'/fund_daily.jpg', width=Inches(7))
        document.add_paragraph(u'投资组合净值如下').paragraph_format.keep_with_next
        document.add_picture(self.plotPath + u'/netValue_daily.jpg', width=Inches(7))
        document.add_paragraph(u'年度收益如下').paragraph_format.keep_with_next
        # document.add_picture(self.plotPath + u'/annualReturns.jpg', width=Inches(7))
        # 保存文件
        document.save('../Output/' + self.strategyNameList[0] + u'-' + self.today + u'.docx')
        print('回测结果写入报告完毕')


'''Test code'''
if __name__== "__main__":
    pass


